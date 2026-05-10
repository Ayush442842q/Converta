from __future__ import annotations

import html
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from markdown import markdown
from markdownify import markdownify as html_to_markdown
from pypdf import PdfReader
from xhtml2pdf import pisa


class ConversionError(RuntimeError):
    """Raised when a document conversion cannot be completed locally."""


@dataclass(frozen=True)
class SupportedFormat:
    extension: str
    name: str


@dataclass(frozen=True)
class ConversionResult:
    source: Path
    destination: Path
    engine: str


class DocumentConverter:
    """Offline document converter with pure-Python, LibreOffice, and Pandoc paths."""

    FORMATS = {
        "md": SupportedFormat("md", "Markdown"),
        "markdown": SupportedFormat("markdown", "Markdown"),
        "txt": SupportedFormat("txt", "Plain text"),
        "html": SupportedFormat("html", "HTML"),
        "htm": SupportedFormat("htm", "HTML"),
        "pdf": SupportedFormat("pdf", "PDF"),
        "docx": SupportedFormat("docx", "Word document"),
        "doc": SupportedFormat("doc", "Legacy Word document"),
        "rtf": SupportedFormat("rtf", "Rich text"),
        "odt": SupportedFormat("odt", "OpenDocument text"),
        "epub": SupportedFormat("epub", "EPUB"),
    }

    PYTHON_TARGETS = {"md", "markdown", "txt", "html", "pdf", "docx"}
    LIBREOFFICE_INPUTS = {"doc", "docx", "odt", "rtf", "txt", "html", "htm", "pdf"}
    LIBREOFFICE_TARGETS = {"pdf", "docx", "doc", "odt", "rtf", "txt", "html"}
    PANDOC_INPUTS = {"md", "markdown", "txt", "html", "htm", "docx", "odt", "rtf", "epub"}
    PANDOC_TARGETS = {"md", "markdown", "txt", "html", "docx", "odt", "rtf", "epub", "pdf"}

    def __init__(self, libreoffice_path: str | None = None, pandoc_path: str | None = None):
        self.libreoffice_path = libreoffice_path or self._find_libreoffice()
        self.pandoc_path = pandoc_path or shutil.which("pandoc")

    @classmethod
    def normalize_extension(cls, value: str) -> str:
        ext = value.lower().strip().lstrip(".")
        if ext == "markdown":
            return "md"
        if ext == "htm":
            return "html"
        return ext

    @classmethod
    def supported_extensions(cls) -> list[str]:
        return sorted(cls.FORMATS)

    def can_convert(self, source_ext: str, target_ext: str) -> bool:
        source = self.normalize_extension(source_ext)
        target = self.normalize_extension(target_ext)
        if source == target:
            return True
        if source not in self.FORMATS or target not in self.FORMATS:
            return False
        if self._can_convert_python(source, target):
            return True
        if self.libreoffice_path and source in self.LIBREOFFICE_INPUTS and target in self.LIBREOFFICE_TARGETS:
            return True
        if self.pandoc_path and source in self.PANDOC_INPUTS and target in self.PANDOC_TARGETS:
            return True
        return False

    def explain_availability(self) -> list[str]:
        notes = ["Built-in Python converters are ready."]
        if self.libreoffice_path:
            notes.append(f"LibreOffice detected: {self.libreoffice_path}")
        else:
            notes.append("LibreOffice was not found. Install LibreOffice for broad Office/PDF conversion.")
        if self.pandoc_path:
            notes.append(f"Pandoc detected: {self.pandoc_path}")
        else:
            notes.append("Pandoc was not found. Install Pandoc for Markdown/HTML/EPUB/DOCX cross-conversion.")
        return notes

    def convert(self, source: Path, destination: Path | None = None, target_ext: str | None = None) -> ConversionResult:
        source = Path(source).resolve()
        if not source.exists():
            raise ConversionError(f"Input file not found: {source}")

        source_ext = self.normalize_extension(source.suffix)
        if not target_ext and not destination:
            raise ConversionError("Choose an output format.")

        if destination:
            destination = Path(destination).resolve()
            target_ext = self.normalize_extension(destination.suffix)
        else:
            target_ext = self.normalize_extension(target_ext or "")
            destination = source.with_suffix(f".{target_ext}").resolve()

        if source_ext == target_ext:
            destination.parent.mkdir(parents=True, exist_ok=True)
            if source != destination:
                shutil.copy2(source, destination)
            return ConversionResult(source, destination, "copy")

        if source_ext not in self.FORMATS:
            raise ConversionError(f"Unsupported input format: .{source_ext}")
        if target_ext not in self.FORMATS:
            raise ConversionError(f"Unsupported output format: .{target_ext}")

        destination.parent.mkdir(parents=True, exist_ok=True)

        if self._can_convert_python(source_ext, target_ext):
            self._convert_python(source, destination, source_ext, target_ext)
            return ConversionResult(source, destination, "python")

        if self.libreoffice_path and source_ext in self.LIBREOFFICE_INPUTS and target_ext in self.LIBREOFFICE_TARGETS:
            self._convert_with_libreoffice(source, destination, target_ext)
            return ConversionResult(source, destination, "libreoffice")

        if self.pandoc_path and source_ext in self.PANDOC_INPUTS and target_ext in self.PANDOC_TARGETS:
            self._convert_with_pandoc(source, destination)
            return ConversionResult(source, destination, "pandoc")

        raise ConversionError(
            f"No offline converter is available for .{source_ext} to .{target_ext}. "
            "Install LibreOffice or Pandoc to unlock more formats."
        )

    def _can_convert_python(self, source_ext: str, target_ext: str) -> bool:
        pairs = {
            ("md", "html"),
            ("md", "txt"),
            ("md", "pdf"),
            ("markdown", "html"),
            ("markdown", "txt"),
            ("markdown", "pdf"),
            ("txt", "html"),
            ("txt", "md"),
            ("txt", "pdf"),
            ("txt", "docx"),
            ("html", "txt"),
            ("html", "md"),
            ("html", "pdf"),
            ("docx", "txt"),
            ("docx", "html"),
            ("docx", "md"),
            ("pdf", "txt"),
        }
        return (source_ext, target_ext) in pairs

    def _convert_python(self, source: Path, destination: Path, source_ext: str, target_ext: str) -> None:
        if source_ext in {"md", "markdown"}:
            text = source.read_text(encoding="utf-8")
            if target_ext == "html":
                destination.write_text(self._markdown_to_full_html(text, source.stem), encoding="utf-8")
            elif target_ext == "txt":
                destination.write_text(text, encoding="utf-8")
            elif target_ext == "pdf":
                self._html_to_pdf(self._markdown_to_full_html(text, source.stem), destination)
            return

        if source_ext == "txt":
            text = source.read_text(encoding="utf-8")
            if target_ext in {"md", "markdown"}:
                destination.write_text(text, encoding="utf-8")
            elif target_ext == "html":
                body = f"<pre>{html.escape(text)}</pre>"
                destination.write_text(self._html_document(body, source.stem), encoding="utf-8")
            elif target_ext == "pdf":
                self._html_to_pdf(self._html_document(f"<pre>{html.escape(text)}</pre>", source.stem), destination)
            elif target_ext == "docx":
                document = Document()
                for paragraph in text.splitlines() or [""]:
                    document.add_paragraph(paragraph)
                document.save(destination)
            return

        if source_ext == "html":
            raw = source.read_text(encoding="utf-8")
            soup = BeautifulSoup(raw, "html.parser")
            if target_ext == "txt":
                destination.write_text(soup.get_text("\n").strip() + "\n", encoding="utf-8")
            elif target_ext in {"md", "markdown"}:
                destination.write_text(html_to_markdown(raw).strip() + "\n", encoding="utf-8")
            elif target_ext == "pdf":
                self._html_to_pdf(raw, destination)
            return

        if source_ext == "docx":
            document = Document(source)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            if target_ext == "txt":
                destination.write_text(text, encoding="utf-8")
            elif target_ext == "html":
                body = "".join(f"<p>{html.escape(paragraph.text)}</p>" for paragraph in document.paragraphs)
                destination.write_text(self._html_document(body, source.stem), encoding="utf-8")
            elif target_ext in {"md", "markdown"}:
                destination.write_text(text, encoding="utf-8")
            return

        if source_ext == "pdf" and target_ext == "txt":
            reader = PdfReader(str(source))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            destination.write_text(text.strip() + "\n", encoding="utf-8")
            return

        raise ConversionError(f"Built-in converter not implemented for .{source_ext} to .{target_ext}.")

    def _convert_with_libreoffice(self, source: Path, destination: Path, target_ext: str) -> None:
        with tempfile.TemporaryDirectory(prefix="converta-lo-") as temp_dir:
            command = [
                self.libreoffice_path,
                "--headless",
                "--convert-to",
                target_ext,
                "--outdir",
                temp_dir,
                str(source),
            ]
            self._run(command)
            produced = next(Path(temp_dir).glob(f"{source.stem}.*"), None)
            if not produced:
                raise ConversionError("LibreOffice did not produce an output file.")
            shutil.move(str(produced), destination)

    def _convert_with_pandoc(self, source: Path, destination: Path) -> None:
        self._run([self.pandoc_path, str(source), "-o", str(destination)])

    def _run(self, command: list[str | None]) -> None:
        try:
            completed = subprocess.run(
                [str(part) for part in command if part],
                check=True,
                capture_output=True,
                text=True,
                timeout=180,
            )
        except FileNotFoundError as exc:
            raise ConversionError(f"Converter executable not found: {command[0]}") from exc
        except subprocess.CalledProcessError as exc:
            message = (exc.stderr or exc.stdout or "The converter failed.").strip()
            raise ConversionError(message) from exc
        except subprocess.TimeoutExpired as exc:
            raise ConversionError("Conversion timed out.") from exc
        if completed.stderr and "Error" in completed.stderr:
            raise ConversionError(completed.stderr.strip())

    def _markdown_to_full_html(self, text: str, title: str) -> str:
        body = markdown(text, extensions=["extra", "tables", "fenced_code", "toc"], output_format="html5")
        return self._html_document(body, title)

    def _html_document(self, body: str, title: str) -> str:
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>{html.escape(title)}</title>
  <style>
    body {{ font-family: Arial, sans-serif; line-height: 1.55; color: #1f2937; margin: 40px; }}
    pre, code {{ font-family: Consolas, monospace; background: #f3f4f6; }}
    pre {{ padding: 16px; white-space: pre-wrap; }}
    table {{ border-collapse: collapse; width: 100%; }}
    th, td {{ border: 1px solid #d1d5db; padding: 6px 8px; }}
  </style>
</head>
<body>
{body}
</body>
</html>
"""

    def _html_to_pdf(self, html_text: str, destination: Path) -> None:
        with destination.open("wb") as file:
            result = pisa.CreatePDF(src=html_text, dest=file)
        if result.err:
            raise ConversionError("PDF generation failed.")

    def _find_libreoffice(self) -> str | None:
        for name in ("soffice", "libreoffice"):
            found = shutil.which(name)
            if found:
                return found

        candidates = [
            Path("C:/Program Files/LibreOffice/program/soffice.exe"),
            Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)
        return None
